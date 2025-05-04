# backend/chatpdf/views.py

from decimal import Decimal
from grpc import Status
from rest_framework.views import APIView
from rest_framework.response import Response
from rest_framework.permissions import IsAuthenticated
from rest_framework_simplejwt.tokens import RefreshToken
from django.contrib.auth import authenticate
import urllib
from .models import PDFDocument, Conversation, Message,User, MessageLimit, Payment
from .serializers import PDFDocumentSerializer, ConversationSerializer, MessageSerializer
from allauth.socialaccount.providers.google.provider import GoogleProvider
from allauth.socialaccount.models import SocialAccount, SocialToken
from allauth.socialaccount.helpers import complete_social_login
from django.contrib.auth import get_user_model
from allauth.socialaccount.providers.oauth2.client import OAuth2Error
from .utils.pdf_utils import read_pdf
from .utils.ai_utils import get_gemini_response, get_llama_response, clear_context, get_deepseek_response, get_qwen_response
from decouple import config
import os
import hashlib
import hmac
from datetime import datetime, timedelta
from pathlib import Path
import re
from django.conf import settings
from django.contrib.auth.models import User  # Thêm dòng này
from django.views.decorators.clickjacking import xframe_options_exempt
from django.http import FileResponse
from docx import Document
from rest_framework import status
import logging
from django.utils import timezone
from allauth.socialaccount.models import SocialApp, SocialAccount
from google.auth.transport import requests
from google.oauth2 import id_token
from django.contrib.auth import get_user_model
from django.core.mail import send_mail
from django.template.loader import render_to_string
from django.utils.http import urlsafe_base64_encode, urlsafe_base64_decode
from django.utils.encoding import force_bytes
from django.contrib.auth.tokens import default_token_generator
import paypalrestsdk
from django.http import HttpResponseRedirect


User = get_user_model()

BASE_DIR = Path(__file__).resolve().parent.parent
User = get_user_model()
logger = logging.getLogger(__name__)

class RegisterView(APIView):
    def post(self, request):
        username = request.data.get('username')
        email = request.data.get('email')
        password = request.data.get('password')
        if User.objects.filter(email=email).exists():
            return Response({"error": "Username exists"}, status=400)
        user = User.objects.create_user(username=username, password=password, email=email)
        return Response({"message": "User created"}, status=201)

class LoginView(APIView):
    def post(self, request):
        email = request.data.get('email')
        password = request.data.get('password')

        if not email or not password:
            return Response({"error": "Email and password are required"}, status=status.HTTP_400_BAD_REQUEST)

        # Tìm user dựa trên email và dùng username để xác thực
        try:
            user = User.objects.get(email=email)
            authenticated_user = authenticate(username=user.username, password=password)
            if authenticated_user:
                refresh = RefreshToken.for_user(authenticated_user)
                return Response({
                    'refresh': str(refresh),
                    'access': str(refresh.access_token),
                }, status=status.HTTP_200_OK)
            return Response({"error": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)
        except User.DoesNotExist:
            return Response({"error": "Invalid credentials"}, status=status.HTTP_400_BAD_REQUEST)

class GoogleLogin(APIView):
    def post(self, request):
        credential = request.data.get('credential')
        logger.info(f"Received credential: {credential}")
        if not credential:
            logger.error("Credential is missing")
            return Response({'error': 'Credential is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            # Xác minh Google ID token
            client_id = config("CLIENT_ID")
            logger.info(f"Using Client ID for verification: {client_id}")
            idinfo = id_token.verify_oauth2_token(credential, requests.Request(), client_id)
            logger.info(f"Verified token info: {idinfo}")

            # Lấy thông tin từ token
            google_id = idinfo['sub']
            email = idinfo['email']
            name = idinfo.get('name', email.split('@')[0])
            picture = idinfo.get('picture', '')  # Lấy URL hình ảnh

            # Tạo hoặc lấy user
            user, created = User.objects.get_or_create(
                username=google_id,
                defaults={
                    'email': email,
                    'first_name': name,
                }
            )
            if created:
                logger.info(f"Created new user: {user.username}")
            else:
                # Cập nhật thông tin nếu user đã tồn tại
                user.email = email
                user.first_name = name
                user.save()
                logger.info(f"Updated user: {user.username}")

            # Tạo hoặc cập nhật SocialAccount (tùy chọn)
            social_account, social_created = SocialAccount.objects.get_or_create(
                user=user,
                provider='google',
                defaults={
                    'uid': google_id,
                    'extra_data': {
                        'email': email,
                        'name': name,
                        'picture': picture,
                    }
                }
            )
            if social_created:
                logger.info(f"Created SocialAccount for user: {user.username}")
            else:
                social_account.extra_data = {
                    'email': email,
                    'name': name,
                    'picture': picture,
                }
                social_account.save()
                logger.info(f"Updated SocialAccount for user: {user.username}")

            # Trả về JWT token
            refresh = RefreshToken.for_user(user)
            return Response({
                'refresh': str(refresh),
                'access': str(refresh.access_token),
                'user': {
                    'id': user.id,
                    'username': user.username,
                    'email': user.email,
                    'name': user.first_name,
                    'picture': picture,  # Trả về picture để frontend dùng
                }
            }, status=status.HTTP_200_OK)

        except ValueError as e:
            logger.error(f"Token verification failed: {str(e)}")
            return Response({'error': f'Invalid Google token: {str(e)}'}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"Unexpected error: {str(e)}")
            return Response({'error': 'Google login failed: ' + str(e)}, status=status.HTTP_400_BAD_REQUEST)

class PasswordResetRequest(APIView):
    def post(self, request):
        email = request.data.get('email')
        logger.info(f"Password reset request for email: {email}")
        
        if not email:
            logger.error("Email is missing")
            return Response({'error': 'Email is required'}, status=status.HTTP_400_BAD_REQUEST)

        try:
            user = User.objects.get(email=email)
            # Tạo token và uid
            token = default_token_generator.make_token(user)
            uid = urlsafe_base64_encode(force_bytes(user.pk))

            # Tạo link reset password
            reset_link = f"http://localhost:5173/reset-password/{uid}/{token}/"  # URL frontend
            logger.info(f"Reset link generated: {reset_link}")

            # Gửi email
            subject = "Reset Your Password"
            message = render_to_string('password_reset_email.html', {
                'user': user,
                'reset_link': reset_link,
            })
            send_mail(
                subject,
                message,
                config('EMAIL_HOST_USER'),
                [email],
                fail_silently=False,
                html_message=message,
            )
            logger.info(f"Password reset email sent to: {email}")
            return Response({'message': 'Password reset link sent to your email'}, status=status.HTTP_200_OK)

        except User.DoesNotExist:
            logger.warning(f"No user found with email: {email}")
            return Response({'error': 'Email not found'}, status=status.HTTP_404_NOT_FOUND)
        except Exception as e:
            logger.error(f"Error sending email: {str(e)}")
            return Response({'error': f'Failed to send email: {str(e)}'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class PasswordResetConfirm(APIView):
    def post(self, request, uidb64, token):
        try:
            uid = urlsafe_base64_decode(uidb64).decode()
            user = User.objects.get(pk=uid)
        except (TypeError, ValueError, OverflowError, User.DoesNotExist):
            user = None

        if user is not None and default_token_generator.check_token(user, token):
            new_password = request.data.get('new_password')
            if not new_password:
                return Response({'error': 'New password is required'}, status=status.HTTP_400_BAD_REQUEST)

            user.set_password(new_password)
            user.save()
            logger.info(f"Password reset successful for user: {user.email}")
            return Response({'message': 'Password reset successful'}, status=status.HTTP_200_OK)
        else:
            logger.error("Invalid reset link or token")
            return Response({'error': 'Invalid reset link or token'}, status=status.HTTP_400_BAD_REQUEST)

class UploadPDFView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        file = request.FILES['file']
        title = request.data.get('title', file.name)
        pdf = PDFDocument.objects.create(user=request.user, file=file, title=title)
        return Response({"message": "PDF uploaded", "id": pdf.id}, status=201)

class ChatView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        user = request.user
        pdf_id = request.data.get('pdf_id')
        message = request.data.get('message', '')
        conversation_id = request.data.get('conversation_id', None)
        model = request.data.get('model', 'gemini')
        title = request.data.get('title')

        # Kiểm tra pdf_id hợp lệ
        try:
            pdf = PDFDocument.objects.get(id=pdf_id, user=user)
        except PDFDocument.DoesNotExist:
            logger.error(f"PDF {pdf_id} not found for user {user.id}")
            return Response({"error": "PDF not found"}, status=status.HTTP_404_NOT_FOUND)

        # Kiểm tra giới hạn tin nhắn nếu không phải tài khoản Plus
        if not user.is_plus:
            today = timezone.now().date()
            message_limit, created = MessageLimit.objects.get_or_create(
                user=user,
                pdf=pdf,
                date=today,
                defaults={'message_count': 0}
            )
            logger.info(f"Message count for user {user.id}, pdf {pdf_id}: {message_limit.message_count}")
            if message_limit.message_count >= 10:
                logger.warning(f"User {user.id} reached message limit for pdf {pdf_id}")
                return Response(
                    {'error': 'Bạn đã đạt giới hạn 10 tin nhắn/ngày cho tài liệu này. Vui lòng nâng cấp lên Plus.'},
                    status=status.HTTP_403_FORBIDDEN
                )
            message_limit.message_count += 1
            message_limit.save()
            logger.info(f"Updated message count: {message_limit.message_count}")

        # Tạo hoặc lấy conversation
        if conversation_id:
            try:
                conversation = Conversation.objects.get(id=conversation_id, user=user)
            except Conversation.DoesNotExist:
                logger.error(f"Conversation {conversation_id} not found for user {user.id}")
                return Response({"error": "Conversation not found"}, status=status.HTTP_404_NOT_FOUND)
        else:
            conversation = Conversation.objects.create(user=user, pdf=pdf, title=title or pdf.title)

        # Lưu tin nhắn user
        Message.objects.create(conversation=conversation, content=message, is_user=True, model_used=model)

        # Đọc PDF chỉ khi tạo conversation mới
        pdf_path = os.path.join(settings.MEDIA_ROOT, pdf.file.name)
        pdf_text = read_pdf(pdf_path) if not conversation_id else None

        # Gọi AI response với try-catch
        try:
            if model == 'gemini':
                ai_reply = get_gemini_response(conversation.id, message, pdf_text)
            elif model == 'deepseek':
                ai_reply = get_deepseek_response(conversation.id, message, pdf_text)
            elif model == 'llama':
                ai_reply = get_llama_response(conversation.id, message, pdf_text)
            elif model == 'qwen':
                ai_reply = get_qwen_response(conversation.id, message, pdf_text)
            else:
                logger.error(f"Invalid model: {model}")
                return Response({"error": "Invalid model"}, status=status.HTTP_400_BAD_REQUEST)
        except Exception as e:
            logger.error(f"AI response error for model {model}: {str(e)}")
            return Response({"error": f"Could not get AI response: {str(e)}"}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

        # Lưu tin nhắn AI
        Message.objects.create(conversation=conversation, content=ai_reply, is_user=False, model_used=model)

        pdf_url = request.build_absolute_uri(f"/api/media/pdfs/{pdf.id}/")
        return Response({"reply": ai_reply, "conversation_id": conversation.id, "pdf_url": pdf_url}, status=status.HTTP_200_OK)

class ConversationHistoryView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        conversations = Conversation.objects.filter(user=request.user)
        serializer = ConversationSerializer(conversations, many=True, context={'request': request})
        return Response(serializer.data)

class ConversationDetailView(APIView):  # Thêm view mới cho CRUD
    permission_classes = [IsAuthenticated]

    def delete(self, request, conversation_id):
        try:
            conversation = Conversation.objects.get(id=conversation_id, user=request.user)
            conversation.delete()
            clear_context(conversation_id)  # Xóa ngữ cảnh khỏi Redis
            return Response({"message": "Conversation deleted"}, status=204)
        except Conversation.DoesNotExist:
            return Response({"error": "Conversation not found"}, status=404)

    def patch(self, request, conversation_id):
        try:
            conversation = Conversation.objects.get(id=conversation_id, user=request.user)
            title = request.data.get('title')
            if title:
                conversation.title = title
                conversation.save()
                return Response({"message": "Conversation updated", "title": title}, status=200)
            return Response({"error": "Title is required"}, status=400)
        except Conversation.DoesNotExist:
            return Response({"error": "Conversation not found"}, status=404)

class MessageHistoryView(APIView):  # Thêm view mới
    permission_classes = [IsAuthenticated]

    def get(self, request, conversation_id):
        try:
            conversation = Conversation.objects.get(id=conversation_id, user=request.user)
            messages = Message.objects.filter(conversation=conversation)
            serializer = MessageSerializer(messages, many=True)
            return Response(serializer.data)
        except Conversation.DoesNotExist:
            return Response({"error": "Conversation not found"}, status=404)

class ServePDFView(APIView):
    permission_classes = [IsAuthenticated]

    @xframe_options_exempt  # Tắt X-Frame-Options cho view này
    def get(self, request, pdf_id):
        try:
            pdf = PDFDocument.objects.get(id=pdf_id, user=request.user)
            return FileResponse(pdf.file.open(), content_type='application/pdf')
        except PDFDocument.DoesNotExist:
            return Response({"error": "PDF not found"}, status=404)

class SummaryMessageByAiView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        conversation_id = request.data.get('conversation_id')
        model = request.data.get('model', 'gemini')

        try:
            conversation = Conversation.objects.get(id=conversation_id, user=request.user)
            messages = Message.objects.filter(conversation=conversation).order_by('created_at')
            serializer = MessageSerializer(messages, many=True)
            message_history = serializer.data
        except Conversation.DoesNotExist:
            return Response({"error": "Conversation not found"}, status=404)

        # Tạo nội dung tin nhắn để tóm tắt
        full_conversation = "\n".join(
            [f"{'User' if msg['is_user'] else 'AI'}: {msg['content']}" for msg in message_history]
        )
        summary_prompt = f"Hãy tóm tắt lại đầy đủ và chi tiết về nội dung trong cuộc hội thoại để tôi làm báo cáo word:\n\n{full_conversation}"

        # Gọi model AI để tóm tắt
        pdf_text = None
        if model == 'gemini':
            summary = get_gemini_response(conversation_id, summary_prompt, pdf_text)
        elif model == 'deepseek':
            summary = get_deepseek_response(conversation_id, summary_prompt, pdf_text)
        elif model == 'llama':
            summary = get_llama_response(conversation_id, summary_prompt, pdf_text)
        elif model == 'gemma':
            summary = get_gemma_response(conversation_id, summary_prompt, pdf_text)
        else:
            return Response({"error": "Invalid model"}, status=400)

        # Tạo nội dung sạch cho file Word
        clean_summary = re.sub(r'[\*\[\]]', '', summary)  # Xóa ký tự rác như *, [, ]

        # Tạo file Word
        doc = Document()
        doc.add_heading(f"Summary of Conversation {conversation_id}", level=1)

        # Chia nội dung thành các dòng và định dạng
        lines = clean_summary.split('\n')
        for line in lines:
            if line.strip():  # Chỉ xử lý các dòng không rỗng
                if ':' in line:  # Nếu dòng có dạng "Tiêu đề: Nội dung"
                    title, content = line.split(':', 1)
                    p = doc.add_paragraph()
                    run = p.add_run(title.strip() + ': ')
                    run.bold = True  # In đậm tiêu đề
                    p.add_run(content.strip())  # Nội dung bình thường
                else:
                    doc.add_paragraph(line.strip())

        # Lưu file Word
        summary_dir = os.path.join(settings.MEDIA_ROOT, 'summaries')
        os.makedirs(summary_dir, exist_ok=True)
        file_path = os.path.join(summary_dir, f"summary_{conversation_id}.docx")
        doc.save(file_path)

        # Tạo URL tải xuống
        summary_url = request.build_absolute_uri(f"/api/media/summaries/summary_{conversation_id}.docx")

        # Lưu nội dung gốc (chưa xóa ký tự rác) vào tin nhắn
        Message.objects.create(
            conversation=conversation,
            content=summary,  # Lưu nội dung gốc từ model AI
            is_user=False,
            model_used=model
        )

        return Response({
            "summary": summary,  # Trả về nội dung gốc cho frontend
            "conversation_id": conversation_id,
            "download_url": summary_url
        })

# Thêm route để tải file
def download_summary(request, filename):
    file_path = os.path.join(settings.MEDIA_ROOT, 'summaries', filename)
    if os.path.exists(file_path):
        return FileResponse(open(file_path, 'rb'), as_attachment=True, filename=filename)
    return Response({"error": "File not found"}, status=404)

class PaypalPaymentView(APIView):
    permission_classes = [IsAuthenticated]

    def post(self, request):
        plan = request.data.get('plan')
        if plan not in ['monthly', 'yearly']:
            return Response({'error': 'Gói không hợp lệ'}, status=status.HTTP_400_BAD_REQUEST)

        amount = Decimal('20.00') if plan == 'monthly' else Decimal('200.00')  # USD
        order_id = f"{request.user.id}_{int(timezone.now().timestamp())}"

        # Tạo bản ghi Payment với trạng thái pending
        payment_record = Payment.objects.create(
            user=request.user,
            amount=amount,
            payment_status='pending',
            payment_method='paypal',
            transaction_id=order_id
        )

        # Cấu hình PayPal SDK
        paypalrestsdk.configure({
            "mode": settings.PAYPAL_MODE,
            "client_id": settings.PAYPAL_CLIENT_ID,
            "client_secret": settings.PAYPAL_SECRET
        })

        # Tạo thanh toán
        payment = paypalrestsdk.Payment({
            "intent": "sale",
            "payer": {"payment_method": "paypal"},
            "redirect_urls": {
                "return_url": f"{settings.PAYPAL_RETURN_URL}api/payment/paypal/callback/?plan={plan}&order_id={order_id}",
                "cancel_url": f"http://localhost:5173/payment/failed?status=canceled&plan={plan}"
            },
            "transactions": [{
                "amount": {
                    "total": f"{amount:.2f}",
                    "currency": "USD"
                },
                "description": f"Nang cap Plus {plan} cho user {request.user.id}"
            }]
        })

        if payment.create():
            approval_url = next(link.href for link in payment.links if link.rel == "approval_url")
            print("Payment URL:", approval_url)
            return Response({'payment_url': approval_url}, status=status.HTTP_200_OK)
        else:
            print("PayPal Error:", payment.error)
            # Cập nhật trạng thái Payment thành failed nếu tạo thanh toán thất bại
            payment_record.payment_status = 'failed'
            payment_record.save()
            return Response({'error': 'Không thể tạo thanh toán PayPal'}, status=status.HTTP_500_INTERNAL_SERVER_ERROR)

class PaypalCallbackView(APIView):
    def get(self, request):
        print("Raw query params:", request.GET)
        print("GET dict:", dict(request.GET))

        payment_id = request.GET.get('paymentId')
        payer_id = request.GET.get('PayerID')
        plan = request.GET.get('plan')
        order_id = request.GET.get('order_id')

        if not (payment_id and payer_id and plan and order_id):
            return Response({
                'error': 'Thiếu tham số thanh toán',
                'debug_info': dict(request.GET)
            }, status=status.HTTP_400_BAD_REQUEST)

        # Cấu hình PayPal SDK
        paypalrestsdk.configure({
            "mode": settings.PAYPAL_MODE,
            "client_id": settings.PAYPAL_CLIENT_ID,
            "client_secret": settings.PAYPAL_SECRET
        })

        # Tìm bản ghi Payment
        try:
            payment_record = Payment.objects.get(transaction_id=order_id)
        except Payment.DoesNotExist:
            return Response({'error': 'Không tìm thấy bản ghi thanh toán'}, status=status.HTTP_400_BAD_REQUEST)

        # Xác nhận thanh toán
        payment = paypalrestsdk.Payment.find(payment_id)
        if payment.execute({"payer_id": payer_id}):
            try:
                user_id = int(order_id.split('_')[0])
                user = User.objects.get(id=user_id)
                amount = Decimal(payment.transactions[0].amount.total)
                plan = 'monthly' if amount == Decimal('20.00') else 'yearly'
                expiry = timezone.now() + timedelta(days=30 if plan == 'monthly' else 365)
                
                # Cập nhật user
                user.is_plus = True
                user.plus_expiry = expiry
                user.save()

                # Cập nhật Payment
                payment_record.payment_status = 'completed'
                payment_record.transaction_id = payment_id
                payment_record.save()

                # Redirect đến trang thanh toán thành công
                redirect_url = f'http://localhost:5173/payment/success?plan={plan}'
                print(f"Redirecting to: {redirect_url}")
                return HttpResponseRedirect(redirect_url)
            except (ValueError, User.DoesNotExist):
                # Cập nhật Payment thành failed nếu user không hợp lệ
                payment_record.payment_status = 'failed'
                payment_record.save()
                return Response({'error': 'User không hợp lệ'}, status=status.HTTP_400_BAD_REQUEST)
        else:
            print("PayPal Error:", payment.error)
            # Cập nhật Payment thành failed
            payment_record.payment_status = 'failed'
            payment_record.save()
            # Redirect đến trang thanh toán thất bại
            redirect_url = f'http://localhost:5173/payment/failed?status=failed&plan={plan}'
            print(f"Redirecting to: {redirect_url}")
            return HttpResponseRedirect(redirect_url)

class UserView(APIView):
    permission_classes = [IsAuthenticated]

    def get(self, request):
        user = request.user
        social_account = SocialAccount.objects.filter(user=user, provider='google').first()
        picture = social_account.extra_data.get('picture', '') if social_account else ''
        return Response({
            'id': user.id,
            'username': user.username,
            'email': user.email,
            'is_plus': user.is_plus,
            'plus_expiry': user.plus_expiry,
            'name': user.first_name,
            'picture': picture,
        }, status=status.HTTP_200_OK)
